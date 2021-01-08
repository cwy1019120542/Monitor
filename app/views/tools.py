import zipfile
import uuid
import shutil
import re
import os
import openpyxl
import docx
import asyncio
from flask import Blueprint, request, current_app
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches
from PyPDF2 import PdfFileReader, PdfFileWriter
from pptx import Presentation
from ..response import success, recognize_error
from ..func_tools import parameter_check, to_xlsx, set_cell_border, return_file, ocr_request, extract_zip
from ..parameter_config import pic_suffix, zip_suffix, pdf_suffix, ppt_suffix, word_suffix, excel_suffix, no_page_suffix

tools_blueprint = Blueprint("tools", __name__)

@tools_blueprint.route("/ocr_waybill", methods=["POST"])
def ocr_waybill():
    files = request.files
    is_success, parameter_check_response = parameter_check(files, [("file", "file", False, zip_suffix)])
    if not is_success:
        return parameter_check_response
    file, file_name = parameter_check_response["file"]
    config = current_app.config
    base_dir = config["OCR_WAYBILL_DIR"]
    file_dir, extract_dir, walk_list = extract_zip(base_dir, file_name, file)
    ocr_url = config["OCR_URL"]
    result_dict = {"letter": {}, "bill": {}}
    letter_re = r"订单.*?号.*?(\d*)\)"
    bill_re = r"订单号(\d*)"
    pic_path_list = []
    for pic_dir, _, pic_name_list in walk_list:
        for pic_name in pic_name_list:
            if not pic_name.endswith(pic_suffix):
                continue
            pic_path = os.path.join(pic_dir, pic_name)
            pic_path_list.append(pic_path)
    # tasks = [ocr_request(ocr_url, i) for i in pic_path_list]
    # try:
    #     loop = asyncio.get_event_loop()
    # except:
    #     loop = asyncio.new_event_loop()
    #     asyncio.set_event_loop(loop)
    #     result_list = loop.run_until_complete(asyncio.gather(*tasks))
    # else:
    #     result_list = asyncio.run_coroutine_threadsafe(asyncio.gather(*tasks), loop)
    result_list = [ocr_request(ocr_url, i) for i in pic_path_list]
    for pic_path, img_data_list in zip(pic_path_list, result_list):
        if not img_data_list:
            continue
        else:
            pic_name = os.path.split(pic_path)[1]
            img_data = "".join(img_data_list)
            letter_waybill_list = re.findall(letter_re, img_data, re.RegexFlag.S)
            bill_waybill_list = re.findall(bill_re, img_data, re.RegexFlag.S)
            # print(letter_waybill_list, bill_waybill_list)
            for letter_waybill in letter_waybill_list:
                if letter_waybill:
                    result_dict["letter"][pic_name] = letter_waybill
            for bill_waybill in bill_waybill_list:
                if bill_waybill:
                    result_dict["bill"][pic_name] = bill_waybill
            # for re_str in re_str_list:
            #     waybill_list = re.findall(re_str, img_data, re.RegexFlag.S)
            #     for waybill in waybill_list:
            #         if waybill:
            #             result_dict[pic_name] = waybill
            #             break
    shutil.rmtree(file_dir)
    return success(result_dict)

@tools_blueprint.route("/excel_join_word", methods=["POST"])
def excel_join_word():
    files = request.files
    is_success, parameter_check_response = parameter_check(files, [("word", "file", False, (".docx", ".DOCX")), ("excel", "file", False, (".XLSX", ".xlsx", ".xls", ".XLS"))])
    if not is_success:
        return parameter_check_response
    word_file, word_name = parameter_check_response["word"]
    excel_file, excel_name = parameter_check_response["excel"]
    base_dir = current_app.config["EXCEL_JOIN_WORD_DIR"]
    file_dir = os.path.join(base_dir, str(uuid.uuid1()))
    if not os.path.exists(file_dir):
        os.makedirs(file_dir)
    word_path = os.path.join(file_dir, word_name)
    excel_path = os.path.join(file_dir, excel_name)
    word_file.save(word_path)
    excel_file.save(excel_path)
    new_excel_path = to_xlsx(excel_path)
    workbook = openpyxl.load_workbook(new_excel_path, data_only=True)
    all_data_group_list = []
    for sheet_name in workbook.sheetnames:
        sheet = workbook[sheet_name]
        data_group_list = [list(i) for i in sheet.values]
        split_data_group_list = []
        len_data_group_list = len(data_group_list)
        for data_group_index, data_group in enumerate(data_group_list, start=1):
            if data_group.count(None) == len(data_group):
                if split_data_group_list:
                    all_data_group_list.append(split_data_group_list)
                    split_data_group_list = []
                continue
            split_data_group_list.append(data_group)
            if data_group_index == len_data_group_list:
                all_data_group_list.append(split_data_group_list)
    workbook.close()
    document = docx.Document(word_path)
    for data_group_list in all_data_group_list:
        clean_data_group_list = []
        data_len = 0
        data_group_list_len = len(data_group_list)
        prefix_data_list = []
        suffix_data_list = []
        for data_group_index, data_group in enumerate(data_group_list):
            clean_data_group = [i for i in data_group if i and "Unnamed:" not in str(i)]
            if not data_len:
                if len(clean_data_group) >= 3:
                    clean_data_group_list.append(clean_data_group)
                    data_len = len(clean_data_group)
                else:
                    prefix_data_list.extend(clean_data_group)
            else:
                if len(clean_data_group) <= 1:
                    next_data_group_index = data_group_index + 1
                    if next_data_group_index <= data_group_list_len - 1:
                        next_data_group = data_group_list[next_data_group_index]
                        clean_next_data_group = [i for i in next_data_group if i and "Unnamed:" not in str(i)]
                        if len(clean_next_data_group) <=1:
                            suffix_data_list.extend(clean_data_group)
                        else:
                            clean_data_group_list.append(data_group[:data_len])
                    else:
                        suffix_data_list.extend(clean_data_group)
                else:
                    clean_data_group_list.append(data_group[:data_len])
        row = len(clean_data_group_list)
        column = data_len
        len_list = [0] * column
        for row_index, row_data in enumerate(clean_data_group_list):
            for column_index, data in enumerate(row_data):
                str_data = str(data)
                if column_index != 0:
                    try:
                        str_data = str(round(float(str_data), 4))
                        clean_data_group_list[row_index][column_index] = str_data
                    except Exception as error:
                        pass
                next_row_index = row_index + 1
                while True:
                    if next_row_index >= row:
                        break
                    if not clean_data_group_list[next_row_index][column_index]:
                        clean_data_group_list[next_row_index][column_index] = str_data
                        next_row_index += 1
                    else:
                        break
                data_len = len(str_data)
                num_list = re.findall(r"\d", str_data)
                num_len = len(num_list)
                char_len = data_len - num_len
                true_len = int(num_len * 0.25) + char_len * 2
                if true_len > len_list[column_index]:
                    len_list[column_index] = true_len
        sum_len = sum(len_list)
        # print(len_list)
        width_rate_list = [i / sum_len for i in len_list]
        # print(width_rate_list)
        section = document.sections[-1]
        new_width, new_height = section.page_height, section.page_width
        if new_width > new_height:
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width = new_width
            section.page_height = new_height
        for prefix_data in prefix_data_list:
            document.add_paragraph(str(prefix_data))
        table = document.add_table(rows=row, cols=column)
        for row_index, row_data_list in enumerate(clean_data_group_list):
            cell_list = table.rows[row_index].cells
            for cell, data in zip(cell_list, row_data_list):
                # paragraph = cell.add_paragraph(str(data if data else ""))
                cell.text = str(data if data else "")
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                # paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                set_cell_border(cell)
        for i in range(column):
            for j in range(row):
                table.cell(j, i).width = Inches(10 * width_rate_list[i])
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for suffix_data in suffix_data_list:
            document.add_paragraph(str(suffix_data))
        document.add_paragraph("\n\n")
    document.save(word_path)
    return return_file(word_path, True, True)

@tools_blueprint.route("/rename_contract", methods=["POST"])
def rename_contract():
    files = request.files
    is_success, parameter_check_response = parameter_check(files, [("file", "file", False, zip_suffix)])
    if not is_success:
        return parameter_check_response
    file, file_name = parameter_check_response["file"]
    config = current_app.config
    base_dir = config["RENAME_CONTRACT_DIR"]
    file_dir, extract_dir, walk_list = extract_zip(base_dir, file_name, file)
    ocr_url = config["OCR_URL"]
    for pic_dir, _, pic_name_list in walk_list:
        for pic_name in pic_name_list:
            pic_path = os.path.join(pic_dir, pic_name)
            if pic_name.endswith(pic_suffix):
                if pic_name.endswith(pdf_suffix):
                    try:
                        pdf_file = PdfFileReader(pic_path, strict=False)
                        page_count = pdf_file.getNumPages()
                        if page_count > 0:
                            first_page = pdf_file.getPage(0)
                            new_pdf_file = PdfFileWriter()
                            new_pdf_file.addPage(first_page)
                            new_pdf_file.write(open(pic_path, 'wb'))
                    except Exception as error:
                        print(f"pdf切割错误 {error}")
                        continue
            else:
                os.remove(pic_path)
    result_dict = {}
    walk_list = os.walk(extract_dir)
    pic_path_list = []
    for pic_dir, _, pic_name_list in walk_list:
        for pic_name in pic_name_list:
            pic_path = os.path.join(pic_dir, pic_name)
            pic_path_list.append(pic_path)
    # tasks = [ocr_request(ocr_url, i) for i in pic_path_list]
    # loop = asyncio.new_event_loop()
    # asyncio.set_event_loop(loop)
    # result_list = loop.run_until_complete(asyncio.gather(*tasks))
    # loop.close()
    result_list = [ocr_request(ocr_url, i) for i in pic_path_list]
    for pic_path, img_data_list in zip(pic_path_list, result_list):
        if not img_data_list:
            continue
        else:
            pic_name = os.path.split(pic_path)[1]
            result_dict[pic_name] = img_data_list
    shutil.rmtree(file_dir)
    return success(result_dict)

@tools_blueprint.route("/page_count", methods=["POST"])
def page_count():
    files = request.files
    is_success, parameter_check_response = parameter_check(files, [("file", "file", False, zip_suffix)])
    if not is_success:
        return parameter_check_response
    file, file_name = parameter_check_response["file"]
    config = current_app.config
    base_dir = config["PAGE_COUNT_DIR"]
    file_dir, extract_dir, walk_list = extract_zip(base_dir, file_name, file)
    page_count_dict = {}
    for pic_dir, _, pic_name_list in walk_list:
        for pic_name in pic_name_list:
            pic_path = os.path.join(pic_dir, pic_name)
            try:
                if pic_name.endswith(pdf_suffix):
                    pdf_file = PdfFileReader(pic_path, strict=False)
                    page_count = pdf_file.getNumPages()
                    page_count_dict[pic_name] = page_count
                elif pic_name.endswith(excel_suffix):
                    new_pic_path = to_xlsx(pic_path)
                    workbook = openpyxl.load_workbook(new_pic_path)
                    page_count = len(workbook.sheetnames)
                    workbook.close()
                    page_count_dict[pic_name] = page_count
                elif pic_name.endswith(word_suffix):
                    document = docx.Document(pic_path)
                    page_count = len(document.sections)
                    page_count_dict[pic_name] = page_count
                elif pic_name.endswith(ppt_suffix):
                    ppt = Presentation(pic_path)
                    page_count = len(ppt.slides)
                    page_count_dict[pic_name] = page_count
                elif pic_name.endswith(no_page_suffix):
                    page_count_dict[pic_name] = 1
                else:
                    page_count_dict[pic_name] = None
            except:
                page_count_dict[pic_name] = None
                print(f"{pic_name} read error")
    shutil.rmtree(file_dir)
    return success(page_count_dict)






